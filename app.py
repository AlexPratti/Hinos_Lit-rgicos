import streamlit as st
from docx import Document
from supabase import create_client
import pandas as pd
import re

# Conexão segura
supabase = create_client(st.secrets["URL_SUPABASE"], st.secrets["KEY_SUPABASE"])

def process_docx(file):
    doc = Document(file)
    data = []
    
    # 1. MAPEAMENTO DO SUMÁRIO
    # Vamos descobrir quais são as seções e hinos lendo o início do arquivo
    secoes_encontradas = []
    hinos_mapeados = []
    
    # Regex para detectar hinos no sumário (Ex: "1. VENHA A NÓS... 10")
    re_hino_sumario = re.compile(r'^(\d+[\.\)].+?)\s?\.+\s?\d+$')
    # Regex para seções no sumário (Texto em CAIXA ALTA seguido de pontinhos)
    re_secao_sumario = re.compile(r'^([A-ZÇÃÕÉÍÓÚ\s]{3,})\s?\.+\s?\d+$')

    # Passo A: Identificar a estrutura pelo Sumário
    current_cat = "OUTROS"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        # Se chegamos na página 10 (onde começam os hinos reais), paramos de ler o sumário
        if text.startswith("ORANTES") and "...." not in text:
            break
            
        m_secao = re_secao_sumario.match(text)
        m_hino = re_hino_sumario.match(text)
        
        if m_secao:
            nome_cat = m_secao.group(1).strip()
            secoes_encontradas.append(nome_cat)
            current_cat = nome_cat
        elif m_hino:
            nome_hino = m_hino.group(1).strip()
            hinos_mapeados.append({"cat": current_cat, "titulo": nome_hino})

    # Passo B: Capturar o conteúdo dos hinos no corpo do texto
    # Criamos um dicionário: Título -> Texto
    conteudos = {}
    hino_atual = None
    buffer_letra = []

    # Lista de títulos para busca exata
    titulos_lista = [h['titulo'] for h in hinos_mapeados]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        # Se a linha for exatamente um dos títulos do sumário
        if text in titulos_lista:
            if hino_atual:
                conteudos[hino_atual] = "\n".join(buffer_letra)
            hino_atual = text
            buffer_letra = []
        elif hino_atual:
            buffer_letra.append(text)

    # Salva o último
    if hino_atual:
        conteudos[hino_atual] = "\n".join(buffer_letra)

    # Passo C: Montar lista final
    for h in hinos_mapeados:
        data.append({
            "n1": h['cat'],
            "n2": h['titulo'],
            "texto": conteudos.get(h['titulo'], "Letra não encontrada no corpo do documento.")
        })
    
    return data

def save_to_db(data):
    # Limpeza profunda
    supabase.table("hinos_conteudos").delete().neq("id", 0).execute()
    supabase.table("hinos_categorias").delete().neq("id", 0).execute()
    
    seen_cats = {}
    for item in data:
        cat_name = item['n1']
        if cat_name not in seen_cats:
            res = supabase.table("hinos_categorias").insert({"nome_nivel1": cat_name}).execute()
            if res.data:
                seen_cats[cat_name] = res.data[0]['id']
        
        if cat_name in seen_cats:
            supabase.table("hinos_conteudos").insert({
                "categoria_id": seen_cats[cat_name],
                "nome_nivel2": item['n2'],
                "texto_completo": item['texto']
            }).execute()

# --- INTERFACE ---
st.set_page_config(page_title="Hinário Litúrgico", layout="wide")

with st.sidebar:
    st.title("⚙️ Painel Admin")
    arquivo = st.file_uploader("Upload do Word (.docx)", type="docx")
    if st.button("🚀 Processar via Sumário"):
        if arquivo:
            with st.spinner("Lendo sumário e extraindo letras..."):
                dados = process_docx(arquivo)
                save_to_db(dados)
                st.success(f"Sucesso! {len(dados)} hinos carregados.")
                st.rerun()

# --- NAVEGAÇÃO ---
try:
    res_cat = supabase.table("hinos_categorias").select("*").order("id").execute()
    if res_cat.data:
        df_cat = pd.DataFrame(res_cat.data)
        
        c1, c2 = st.columns([1, 2])
        
        with c1:
            cat_sel = st.selectbox("📌 1. Selecione a Seção:", df_cat['nome_nivel1'])
            cat_id = int(df_cat[df_cat['nome_nivel1'] == cat_sel]['id'].iloc[0])
            
            busca = st.text_input("🔍 2. Busca rápida:")
            
            hinos_db = supabase.table("hinos_conteudos").select("*").eq("categoria_id", cat_id).order("id").execute().data
            
            if hinos_db:
                if busca:
                    hinos_db = [h for h in hinos_db if busca.lower() in h['nome_nivel2'].lower()]
                
                titulos = [h['nome_nivel2'] for h in hinos_db]
                if titulos:
                    escolha = st.radio("📑 3. Escolha o hino:", titulos)
                    info = next(h for h in hinos_db if h['nome_nivel2'] == escolha)
                    
                    with c2:
                        st.subheader(info['nome_nivel2'])
                        st.divider()
                        st.text(info['texto_completo'])
                else:
                    st.warning("Nenhum hino encontrado.")
    else:
        st.info("💡 Banco de dados vazio. Faça o upload no menu lateral.")
except Exception as e:
    st.error(f"Erro: {e}")
